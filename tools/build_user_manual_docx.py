from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SHOT = ROOT / "manual_shots_full"
OUT = ROOT / "USER_MANUAL_NSD_LENGKAP_BAB_SCREENSHOT.docx"

ACCENT = RGBColor(198, 43, 57)
GREEN = RGBColor(45, 134, 105)
INK = RGBColor(45, 29, 30)
MUTED = RGBColor(112, 92, 92)
LIGHT = "FFF7F5"
HEADER = "E8EEF5"

SCREENS = {
    "user_home": SHOT / "app-user-01-home.png",
    "user_campaign": SHOT / "app-user-02-campaign.png",
    "user_counseling": SHOT / "app-user-03-konseling.png",
    "user_transparency": SHOT / "app-user-04-transparansi.png",
    "user_login": SHOT / "app-user-05-login.png",
    "counsel_login": SHOT / "app-counseling-01-login.png",
    "admin_login": SHOT / "web-admin-01-login.png",
    "admin_after": SHOT / "web-admin-02-after-login.png",
    "admin_dashboard": SHOT / "web-admin-03-dashboard.png",
}


def set_run_font(run, name: str = "Calibri") -> None:
    run.font.name = name
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:ascii"), name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(cell, color: str = "D9C8C5") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def keep_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "User Manual NSD - Nusantara Spiritual Donation"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("NSD User Manual | APK User | APK Counseling | Web Admin")
    set_run_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    keep_next(p)
    for run in p.runs:
        set_run_font(run)
        if level == 1:
            run.font.color.rgb = ACCENT
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = GREEN
            run.font.size = Pt(13)
        else:
            run.font.color.rgb = INK
            run.font.size = Pt(12)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK


def steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK


def callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width(table, [6.4])
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    border(cell, "E5CBC6")
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run_font(run)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(10.5)
    p2 = cell.add_paragraph(body)
    for r in p2.runs:
        set_run_font(r)
        r.font.size = Pt(10)
    doc.add_paragraph()


def matrix(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, HEADER)
        border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run)
                run.font.bold = True
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run)
                    run.font.size = Pt(9.2)
    doc.add_paragraph()


def figure(doc: Document, image: Path, caption: str, explanation: str, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    exp = doc.add_paragraph()
    exp.paragraph_format.left_indent = Inches(0.35)
    exp.paragraph_format.right_indent = Inches(0.35)
    run = exp.add_run(explanation)
    set_run_font(run)
    run.font.size = Pt(9.6)
    run.font.color.rgb = MUTED
    doc.add_paragraph()


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("USER MANUAL")
    set_run_font(run)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Nusantara Spiritual Donation (NSD)")
    set_run_font(run)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Panduan lengkap aplikasi user, aplikasi konseling, web admin, Firebase, Railway, dan alur operasional"
    )
    set_run_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    matrix(
        doc,
        ["Metadata", "Isi"],
        [
            ("Nama Sistem", "NSD - Nusantara Spiritual Donation"),
            ("Platform", "APK User, APK Counseling, Web Admin"),
            ("Stack", "Flutter, Node.js API, Firebase, Railway, Firebase Hosting"),
            ("Tanggal Dokumen", datetime.now().strftime("%d %B %Y")),
        ],
        [1.7, 4.6],
    )
    callout(
        doc,
        "Catatan penting",
        "Screenshot dalam dokumen ini diambil ulang dari build aplikasi yang berjalan di workspace, bukan dari gambar referensi yang dikirimkan. Tampilan dapat berubah apabila warna, copywriting, atau data seed diubah pada source code.",
    )
    doc.add_page_break()


def build() -> None:
    for name, path in SCREENS.items():
        if not path.exists():
            raise FileNotFoundError(f"Screenshot tidak ditemukan: {name}: {path}")

    doc = Document()
    configure(doc)
    title_page(doc)

    heading(doc, "Daftar Isi", 1)
    bullets(
        doc,
        [
            "BAB 1 Pendahuluan",
            "BAB 2 Ringkasan Sistem dan Role Pengguna",
            "BAB 3 Arsitektur Aplikasi dan Deployment",
            "BAB 4 Persiapan Sebelum Menggunakan Sistem",
            "BAB 5 Panduan Aplikasi User / Donatur",
            "BAB 6 Panduan Aplikasi Konseling / Konselor",
            "BAB 7 Panduan Web Admin Panel",
            "BAB 8 Data, Firebase, dan Railway",
            "BAB 9 SOP Operasional Harian",
            "BAB 10 Troubleshooting",
            "Lampiran A Akun Demo dan Struktur Data",
            "Lampiran B Command Build dan Deploy",
        ],
    )
    callout(
        doc,
        "Cara membaca dokumen",
        "BAB 1 sampai BAB 4 menjelaskan konteks dan setup. BAB 5 sampai BAB 7 dipakai sebagai panduan layar per layar. BAB 8 sampai lampiran dipakai untuk teknis operasional, deployment, dan troubleshooting.",
    )

    doc.add_page_break()
    heading(doc, "BAB 1 Pendahuluan", 1)
    para(
        doc,
        "Nusantara Spiritual Donation (NSD) adalah platform donasi digital yang menghubungkan donatur, pemohon bantuan, konselor, operator, admin, dan super admin dalam satu alur kerja. Sistem ini memisahkan aplikasi mobile untuk pengguna umum dari web admin panel agar pengalaman donatur tetap sederhana, sementara pengelolaan data tetap terkontrol.",
    )
    para(
        doc,
        "Tujuan utama manual ini adalah menjelaskan cara menggunakan aplikasi, memahami fungsi setiap halaman, menjalankan proses admin, dan menangani masalah umum yang muncul saat build, deploy, atau penggunaan harian.",
    )
    heading(doc, "1.1 Ruang Lingkup", 2)
    bullets(
        doc,
        [
            "Panduan penggunaan APK User untuk donatur dan pemohon bantuan.",
            "Panduan penggunaan APK Counseling untuk konselor atau pendamping.",
            "Panduan penggunaan Web Admin Panel untuk admin, operator, super admin, dan konselor.",
            "Ringkasan setup Firebase, Firestore, Railway, dan Firebase Hosting.",
            "Troubleshooting layar putih, login gagal, data tidak muncul, APK tidak terpasang, dan deploy tidak berjalan.",
        ],
    )
    heading(doc, "1.2 Platform yang Dibuat", 2)
    matrix(
        doc,
        ["Platform", "Target Pengguna", "Fungsi Utama"],
        [
            ("APK User", "Donatur dan pemohon bantuan", "Melihat campaign, berdonasi, melihat transparansi, dan mengakses konseling publik."),
            ("APK Counseling", "Konselor", "Masuk sebagai konselor, melihat kebutuhan pendampingan, dan mengelola interaksi konseling."),
            ("Web Admin Panel", "Admin, operator, super admin, konselor", "Mengelola campaign, bantuan, pengguna, audit, transparansi, dan operasional sistem."),
        ],
        [1.4, 1.8, 3.0],
    )

    heading(doc, "BAB 2 Ringkasan Sistem dan Role Pengguna", 1)
    para(
        doc,
        "NSD menggunakan konsep role agar setiap pengguna hanya melihat fitur yang sesuai kewenangannya. Pemisahan role ini penting karena sistem menangani data donasi, data bantuan, percakapan konseling, dan laporan penyaluran dana.",
    )
    heading(doc, "2.1 Daftar Role", 2)
    matrix(
        doc,
        ["Role", "Akses", "Catatan"],
        [
            ("Donatur", "Melihat campaign, donasi, riwayat, transparansi, konseling publik.", "Role default untuk pengguna yang ingin berdonasi."),
            ("Pemohon", "Mengajukan bantuan, melengkapi data, melihat status pengajuan.", "Dipakai oleh pihak yang membutuhkan bantuan."),
            ("Konselor", "Mendampingi sesi konseling dan membaca konteks pemohon.", "Memakai APK Counseling atau web panel sesuai alur."),
            ("Operator", "Memverifikasi campaign, donasi, dan pengajuan.", "Role kerja operasional harian."),
            ("Admin", "Mengelola data utama, status campaign, dan validasi.", "Role pengelola sistem."),
            ("Super Admin", "Akses tertinggi termasuk audit dan konfigurasi.", "Dipakai untuk pengawasan penuh."),
        ],
        [1.25, 3.2, 2.1],
    )
    heading(doc, "2.2 Alur Umum", 2)
    steps(
        doc,
        [
            "User membuka APK User dan melihat daftar campaign aktif.",
            "User memilih campaign, membaca detail, lalu melakukan donasi atau melihat laporan transparansi.",
            "Jika membutuhkan pendampingan, user membuka menu Konseling dan memilih konselor.",
            "Konselor membuka APK Counseling untuk menangani kebutuhan pendampingan.",
            "Admin membuka Web Admin Panel untuk memonitor campaign, user, pengajuan bantuan, dan audit.",
        ],
    )

    heading(doc, "BAB 3 Arsitektur Aplikasi dan Deployment", 1)
    para(
        doc,
        "Aplikasi dibangun dengan Flutter untuk mobile dan web. Backend API menggunakan Node.js dan dapat berjalan di Railway. Firebase dipakai untuk hosting web, autentikasi, dan/atau Firestore sesuai konfigurasi project. Pemisahan target dilakukan lewat entrypoint Flutter dan product flavor Android.",
    )
    heading(doc, "3.1 Komponen Teknis", 2)
    matrix(
        doc,
        ["Komponen", "Lokasi/Target", "Keterangan"],
        [
            ("Flutter Client", "client/lib", "Berisi UI app user, app counseling, dan web admin."),
            ("Entrypoint User", "client/lib/main_user.dart", "Dipakai untuk APK User dan screenshot app user."),
            ("Entrypoint Counseling", "client/lib/main_counseling.dart", "Dipakai untuk APK Counseling."),
            ("Entrypoint Admin", "client/lib/main.dart", "Dipakai untuk Web Admin Panel."),
            ("Node API", "server/src", "API autentikasi, campaign, donasi, konseling, dan dashboard."),
            ("Firebase Hosting", "client/build/web", "Hosting web admin atau build web sesuai target deploy."),
            ("Railway", "@nsd/server", "Tempat backend API production berjalan."),
        ],
        [1.55, 2.25, 2.6],
    )
    heading(doc, "3.2 Pemisahan APK dan Web", 2)
    para(
        doc,
        "Sistem tidak menjadikan web Railway sebagai app user. Web admin tetap ditujukan untuk pengelola, sedangkan user dan konselor menggunakan APK terpisah di HP. Pada Android, product flavor membuat package berbeda sehingga dua APK dapat dipasang bersamaan.",
    )
    bullets(
        doc,
        [
            "APK User memakai package id.nsd.nsd_app.user dan label NSD User.",
            "APK Counseling memakai package id.nsd.nsd_app.counseling dan label NSD Counseling.",
            "Web Admin memakai entrypoint main.dart dan dapat dideploy ke Firebase Hosting atau diserve oleh backend sesuai kebutuhan.",
        ],
    )

    heading(doc, "BAB 4 Persiapan Sebelum Menggunakan Sistem", 1)
    para(
        doc,
        "Sebelum memakai sistem, pastikan backend, Firebase, dan APK sudah disiapkan. Pada mode demo, data seed menyediakan akun admin, donatur, dan konselor agar pengujian bisa dilakukan tanpa membuat semua data dari awal.",
    )
    heading(doc, "4.1 Kebutuhan Minimum", 2)
    bullets(
        doc,
        [
            "Flutter SDK dan Android SDK untuk build APK.",
            "Node.js 20 atau lebih baru untuk menjalankan server.",
            "Firebase CLI untuk deploy hosting.",
            "Akun Firebase project nsd-donasi dan akses ke Firestore/Hosting.",
            "Railway project untuk backend API production.",
            "HP Android dengan USB debugging aktif jika ingin install APK langsung dari komputer.",
        ],
    )
    heading(doc, "4.2 Pemeriksaan Awal", 2)
    steps(
        doc,
        [
            "Jalankan flutter devices untuk memastikan HP terbaca.",
            "Jalankan npm run typecheck untuk memastikan server dan Flutter analyzer bersih.",
            "Build APK user dan counseling sesuai perintah pada Lampiran B.",
            "Pastikan API Railway atau API lokal dapat diakses melalui endpoint /api/health.",
            "Pastikan Firebase Hosting mengarah ke folder build web yang benar.",
        ],
    )

    heading(doc, "BAB 5 Panduan Aplikasi User / Donatur", 1)
    para(
        doc,
        "APK User adalah aplikasi utama untuk donatur dan pemohon bantuan. Desainnya dibuat mobile-first agar pengguna bisa cepat membaca campaign, memilih aksi, dan memantau laporan dana tanpa masuk ke panel admin.",
    )
    heading(doc, "5.1 Halaman Beranda", 2)
    figure(
        doc,
        SCREENS["user_home"],
        "Gambar 5.1 - Beranda APK User",
        "Halaman Beranda menampilkan identitas NSD, pesan utama bantuan cepat, tombol Lihat Campaign, tombol Ajukan Bantuan, indikator verifikasi, dan kartu campaign awal. Pengguna baru biasanya mulai dari halaman ini untuk memahami fungsi aplikasi.",
        3.05,
    )
    bullets(
        doc,
        [
            "Tombol Lihat Campaign membawa user ke daftar campaign aktif.",
            "Tombol Ajukan Bantuan dipakai pemohon untuk mulai proses bantuan.",
            "Ikon profil di kanan atas membuka halaman masuk atau dashboard user.",
            "Bottom navigation menyediakan akses cepat ke Beranda, Campaign, Konseling, dan Transparansi.",
        ],
    )
    heading(doc, "5.2 Halaman Campaign", 2)
    figure(
        doc,
        SCREENS["user_campaign"],
        "Gambar 5.2 - Daftar Campaign Terverifikasi",
        "Halaman Campaign menampilkan daftar kebutuhan bantuan yang sudah melewati proses verifikasi. User bisa mencari campaign, memilih kategori, dan membuka campaign tertentu untuk membaca detail sebelum donasi.",
        3.05,
    )
    bullets(
        doc,
        [
            "Kolom pencarian membantu menemukan campaign berdasarkan nama, lokasi, atau kategori.",
            "Filter kategori membantu user memilah bantuan bencana, alam, pendidikan, atau kebutuhan lain.",
            "Kartu campaign menampilkan ringkasan lokasi, judul, deskripsi pendek, dan status tampilan awal.",
            "Data campaign diambil dari API atau data seed sehingga daftar tetap muncul untuk demo.",
        ],
    )
    heading(doc, "5.3 Halaman Konseling Publik", 2)
    figure(
        doc,
        SCREENS["user_counseling"],
        "Gambar 5.3 - Menu Konseling di APK User",
        "Menu Konseling menampilkan daftar konselor yang dapat dipilih user. Fungsi ini disiapkan untuk pendampingan spiritual, keluarga, pemulihan emosi, dan dukungan umum bagi penerima bantuan atau user yang membutuhkan.",
        3.05,
    )
    bullets(
        doc,
        [
            "User melihat nama konselor dan bidang pendampingannya.",
            "Tombol Mulai chat menjadi titik awal percakapan konseling.",
            "Apabila API tidak tersedia, aplikasi tetap menampilkan seed counselor agar halaman tidak kosong.",
            "Konseling publik di APK User berbeda dengan APK Counseling yang digunakan oleh pihak konselor.",
        ],
    )
    heading(doc, "5.4 Halaman Transparansi", 2)
    figure(
        doc,
        SCREENS["user_transparency"],
        "Gambar 5.4 - Laporan Transparansi",
        "Halaman Transparansi memperlihatkan angka dana terkumpul, dana disalurkan, saldo program, dan aktivitas donasi. Halaman ini menjadi bukti bahwa donasi tidak hanya diterima tetapi juga dilaporkan secara terbuka.",
        3.05,
    )
    bullets(
        doc,
        [
            "Dana terkumpul menunjukkan total nilai donasi yang sudah masuk.",
            "Dana disalurkan menunjukkan uang yang sudah dipakai untuk program.",
            "Saldo program menunjukkan sisa dana yang belum disalurkan.",
            "Aktivitas donasi membantu user melihat dinamika donasi terakhir.",
        ],
    )
    heading(doc, "5.5 Login User", 2)
    figure(
        doc,
        SCREENS["user_login"],
        "Gambar 5.5 - Login User",
        "Form login dipakai user untuk masuk sebagai donatur atau pemohon. Pada mode demo, akun donatur@nsd.id dengan password Demo1234 dapat dipakai untuk pengujian alur umum.",
        3.05,
    )
    steps(
        doc,
        [
            "Tekan ikon profil pada pojok kanan atas.",
            "Masukkan email dan password akun.",
            "Tekan Masuk.",
            "Jika belum punya akun, gunakan tautan Daftar untuk registrasi.",
            "Setelah berhasil, user dapat membuka dashboard atau melanjutkan donasi.",
        ],
    )

    heading(doc, "BAB 6 Panduan Aplikasi Konseling / Konselor", 1)
    para(
        doc,
        "APK Counseling adalah aplikasi terpisah untuk konselor. Tujuannya agar konselor tidak memakai tampilan donatur biasa, tetapi masuk ke alur pendampingan dan dashboard yang relevan dengan pekerjaan konseling.",
    )
    heading(doc, "6.1 Login Konselor", 2)
    figure(
        doc,
        SCREENS["counsel_login"],
        "Gambar 6.1 - Login APK Counseling",
        "APK Counseling langsung menampilkan form login. Konselor memakai akun dengan role konselor, misalnya konselor@nsd.id pada data demo. Setelah login sukses, aplikasi menampilkan dashboard atau daftar sesi yang perlu ditangani.",
        3.05,
    )
    steps(
        doc,
        [
            "Buka aplikasi NSD Counseling di HP.",
            "Masukkan akun konselor yang sudah dibuat admin.",
            "Tekan Masuk.",
            "Jika login gagal, pastikan akun memiliki role konselor di backend atau Firestore.",
            "Jika aplikasi berhenti di splash, rebuild APK dengan entrypoint main_counseling.dart dan pastikan startup session tidak menahan runApp.",
        ],
    )
    heading(doc, "6.2 Fungsi Konselor", 2)
    bullets(
        doc,
        [
            "Melihat daftar kebutuhan pendampingan dari user atau pemohon.",
            "Membaca konteks awal sebelum merespons percakapan.",
            "Memberikan arahan awal atau dukungan emosional sesuai kebutuhan.",
            "Membantu admin/operator menilai apakah pengajuan bantuan perlu tindak lanjut.",
            "Menjaga percakapan tetap aman, sopan, dan sesuai etika pendampingan.",
        ],
    )
    heading(doc, "6.3 Catatan Operasional Konseling", 2)
    callout(
        doc,
        "Privasi dan keamanan",
        "Data percakapan konseling sebaiknya diperlakukan sebagai data sensitif. Konselor tidak boleh membagikan isi percakapan ke pihak luar selain admin/operator yang berwenang menangani kasus.",
    )

    heading(doc, "BAB 7 Panduan Web Admin Panel", 1)
    para(
        doc,
        "Web Admin Panel adalah dashboard pengelola. Halaman ini bukan untuk donatur umum. Admin panel dipakai untuk melihat data sistem, memverifikasi campaign, meninjau pengajuan bantuan, mengatur transparansi, dan memantau aktivitas operasional.",
    )
    heading(doc, "7.1 Halaman Login Admin", 2)
    figure(
        doc,
        SCREENS["admin_login"],
        "Gambar 7.1 - Web Admin Login",
        "Halaman login admin menampilkan kartu fungsi utama seperti Kelola Campaign, Konseling, Pengajuan Bantuan, serta Audit & Transparansi. Form login sudah mengarahkan user ke akun demo admin apabila data seed aktif.",
        6.25,
    )
    steps(
        doc,
        [
            "Buka URL admin web.",
            "Pastikan backend API berjalan dan dapat diakses dari browser.",
            "Masukkan email admin dan password.",
            "Tekan Masuk ke Panel.",
            "Jika muncul error, periksa API_URL, CORS CLIENT_ORIGIN, dan data akun admin.",
        ],
    )
    heading(doc, "7.2 Status Setelah Login Admin", 2)
    figure(
        doc,
        SCREENS["admin_after"],
        "Gambar 7.2 - Admin Sudah Terautentikasi",
        "Setelah login berhasil, admin panel menampilkan kartu profil admin dan tombol Masuk Panel. Ini menunjukkan token autentikasi sudah diterima dan session user sudah aktif.",
        6.25,
    )
    bullets(
        doc,
        [
            "Tombol Keluar menghapus session dan mengembalikan admin ke halaman login.",
            "Tombol Masuk Panel membuka dashboard operasional.",
            "Nama dan role user ditampilkan agar admin tahu akun yang sedang aktif.",
            "Session memakai token dari backend atau Firebase sesuai konfigurasi build.",
        ],
    )
    heading(doc, "7.3 Dashboard Admin", 2)
    figure(
        doc,
        SCREENS["admin_dashboard"],
        "Gambar 7.3 - Dashboard Admin Panel",
        "Dashboard admin menampilkan navigasi samping, ringkasan metrik, daftar campaign, dan aktivitas terbaru. Dari halaman ini admin dapat berpindah ke pengajuan, pengguna, donasi, bantuan, audit, dan transparansi.",
        6.25,
    )
    bullets(
        doc,
        [
            "Sidebar kiri dipakai untuk berpindah modul.",
            "Kartu ringkasan membantu admin melihat kondisi sistem secara cepat.",
            "Daftar campaign menampilkan progress tiap campaign.",
            "Aktivitas terbaru membantu audit ringan terhadap perubahan data.",
        ],
    )
    heading(doc, "7.4 Modul Admin yang Disarankan", 2)
    matrix(
        doc,
        ["Modul", "Pengguna", "Fungsi"],
        [
            ("Ringkasan", "Admin/operator", "Melihat metrik utama, campaign aktif, dan aktivitas terbaru."),
            ("Pengajuan Bantuan", "Operator/admin/konselor", "Meninjau dokumen pemohon dan status rekomendasi."),
            ("Kelola Campaign", "Admin/operator", "Membuat, memverifikasi, mengaktifkan, atau menutup campaign."),
            ("Data Donasi", "Admin/operator", "Memantau pembayaran, status, nominal, dan riwayat donasi."),
            ("Pengguna", "Admin/super admin", "Mengelola akun dan role pengguna."),
            ("Audit & Transparansi", "Admin/super admin", "Memantau log perubahan dan laporan publik."),
        ],
        [1.5, 1.8, 3.1],
    )

    heading(doc, "BAB 8 Data, Firebase, dan Railway", 1)
    para(
        doc,
        "Data NSD dapat berjalan melalui backend API Node.js dan Firebase. Untuk production, backend berada di Railway, sementara Firebase digunakan untuk hosting web dan layanan Firebase lain sesuai konfigurasi project. Kunci service account Firebase tidak boleh dimasukkan ke GitHub.",
    )
    heading(doc, "8.1 Firestore Collection yang Disarankan", 2)
    matrix(
        doc,
        ["Collection", "Isi Data", "Dipakai Oleh"],
        [
            ("users", "Profil user, email, role, status verifikasi.", "Auth, admin panel, dashboard."),
            ("campaigns", "Judul campaign, kategori, lokasi, target dana, progress.", "APK User, admin panel."),
            ("donations", "Nominal donasi, status, metode pembayaran, waktu.", "Transparansi, admin donasi."),
            ("aidApplications", "Pengajuan bantuan, dokumen, status, catatan.", "Pemohon, operator, konselor."),
            ("counselingSessions", "Sesi konseling, user, konselor, status.", "APK Counseling dan admin."),
            ("auditLogs", "Log perubahan data dan tindakan admin.", "Super admin dan audit."),
        ],
        [1.7, 3.1, 1.8],
    )
    heading(doc, "8.2 Railway Environment Variables", 2)
    matrix(
        doc,
        ["Variable", "Contoh Isi", "Keterangan"],
        [
            ("PORT", "4000 atau otomatis dari Railway", "Port server API."),
            ("CLIENT_ORIGIN", "https://nsd-donasi.web.app", "Origin frontend yang boleh mengakses API."),
            ("JWT_SECRET", "secret panjang minimal 32 karakter", "Kunci tanda tangan token."),
            ("FIREBASE_PROJECT_ID", "nsd-donasi", "ID project Firebase."),
            ("FIREBASE_SERVICE_ACCOUNT_JSON", "{...}", "Service account JSON lengkap. Jangan dipush ke GitHub."),
            ("DATA_FILE", "server/data/nsd.json", "Opsional untuk mode file JSON lokal."),
        ],
        [2.2, 2.2, 2.0],
    )
    heading(doc, "8.3 Firebase Hosting", 2)
    steps(
        doc,
        [
            "Build Flutter web sesuai target yang ingin dideploy.",
            "Pastikan firebase.json mengarah ke folder client/build/web.",
            "Login Firebase CLI memakai akun yang punya akses project nsd-donasi.",
            "Jalankan firebase deploy --only hosting dari root repo.",
            "Cek URL hosting dan refresh browser dengan cache bersih jika halaman masih putih.",
        ],
    )

    heading(doc, "BAB 9 SOP Operasional Harian", 1)
    para(
        doc,
        "SOP ini dipakai agar admin, operator, dan konselor punya pola kerja yang konsisten. Sistem donasi harus menjaga kecepatan respons sekaligus akurasi verifikasi.",
    )
    heading(doc, "9.1 SOP Campaign Baru", 2)
    steps(
        doc,
        [
            "Operator menerima data kebutuhan bantuan atau laporan kejadian.",
            "Operator memeriksa kelengkapan lokasi, kategori, narasi, target dana, dan bukti pendukung.",
            "Admin membuat campaign atau mengubah status draft menjadi verifikasi.",
            "Konselor dapat membantu membaca konteks kemanusiaan jika pengajuan berasal dari pemohon.",
            "Admin mengaktifkan campaign setelah data dianggap valid.",
            "Setelah campaign berjalan, admin memantau donasi dan laporan penyaluran.",
        ],
    )
    heading(doc, "9.2 SOP Donasi", 2)
    steps(
        doc,
        [
            "Donatur memilih campaign di APK User.",
            "Donatur membaca ringkasan dan transparansi campaign.",
            "Donatur memilih nominal dan metode pembayaran.",
            "Sistem mencatat status donasi.",
            "Admin/operator memantau status transaksi pada panel.",
            "Data donasi masuk ke laporan transparansi.",
        ],
    )
    heading(doc, "9.3 SOP Konseling", 2)
    steps(
        doc,
        [
            "User membuka menu Konseling dan memilih konselor.",
            "Konselor masuk ke APK Counseling.",
            "Konselor membaca konteks awal user.",
            "Konselor memberi respons awal yang aman dan tidak menghakimi.",
            "Jika kasus terkait bantuan, konselor memberi catatan untuk operator/admin.",
            "Admin mengarsipkan aktivitas sesuai kebutuhan audit.",
        ],
    )

    heading(doc, "BAB 10 Troubleshooting", 1)
    para(
        doc,
        "Bagian ini menjawab masalah yang sering muncul saat membangun, menjalankan, atau deploy project NSD. Gunakan bagian ini sebelum mengubah source code besar-besaran.",
    )
    heading(doc, "10.1 Web Putih atau Blank", 2)
    bullets(
        doc,
        [
            "Pastikan build web memakai entrypoint yang benar: main.dart untuk admin, main_user.dart untuk user, main_counseling.dart untuk counseling.",
            "Buka Developer Console di browser dan cek error main.dart.js.",
            "Pastikan file build/web sudah dideploy ke Firebase Hosting.",
            "Jika memakai Flutter web, lakukan hard refresh atau clear cache karena service worker bisa menyimpan build lama.",
            "Pastikan runApp tidak tertahan oleh proses async Firebase atau session restore.",
        ],
    )
    heading(doc, "10.2 APK Counseling Freeze di Splash", 2)
    bullets(
        doc,
        [
            "Pastikan APK dibuild dengan flutter build apk --release --flavor counseling -t lib/main_counseling.dart.",
            "Pastikan session restore tidak menahan tampilan pertama terlalu lama.",
            "Pastikan API_URL untuk Android mengarah ke server yang bisa diakses HP fisik, bukan 10.0.2.2 kecuali menggunakan emulator.",
            "Ambil log dengan adb logcat jika splash tidak hilang setelah 10 detik.",
            "Install ulang APK setelah rebuild agar package lama tidak tertinggal.",
        ],
    )
    heading(doc, "10.3 Firebase Deploy Gagal Project Tidak Ditemukan", 2)
    bullets(
        doc,
        [
            "Jalankan firebase logout lalu firebase login dengan akun yang memiliki akses project.",
            "Cek firebase projects:list dan pastikan nsd-donasi muncul.",
            "Pastikan .firebaserc berisi project ID yang benar.",
            "Jalankan firebase deploy --only hosting dari root repo yang memiliki firebase.json.",
        ],
    )
    heading(doc, "10.4 Railway Not Found atau Blank", 2)
    bullets(
        doc,
        [
            "Cek service Railway active dan deploy logs tidak error.",
            "Pastikan domain Railway mengarah ke service backend yang benar.",
            "Cek endpoint /api/health untuk memastikan API hidup.",
            "Jika web admin dipisah ke Firebase Hosting, Railway tidak wajib menampilkan app user.",
            "Pastikan CLIENT_ORIGIN di Railway sesuai domain frontend production.",
        ],
    )
    heading(doc, "10.5 Login Admin Gagal", 2)
    bullets(
        doc,
        [
            "Pastikan API_URL frontend mengarah ke API yang benar.",
            "Pastikan CORS CLIENT_ORIGIN mengizinkan domain web admin.",
            "Pastikan akun admin ada di data seed atau Firestore.",
            "Pastikan JWT_SECRET tidak kosong.",
            "Cek Network tab di browser untuk melihat status response /auth/login.",
        ],
    )

    heading(doc, "Lampiran A Akun Demo dan Struktur Data", 1)
    para(doc, "Akun demo berikut dipakai untuk pengujian lokal atau mode seed. Jangan menggunakan password demo untuk production.")
    matrix(
        doc,
        ["Akun", "Email", "Password", "Role"],
        [
            ("Donatur Demo", "donatur@nsd.id", "Demo1234", "donatur"),
            ("Konselor Demo", "konselor@nsd.id", "Demo1234", "konselor"),
            ("Admin Demo", "admin@nsd.id", "Demo1234", "admin"),
            ("Super Admin Demo", "superadmin@nsd.id", "Demo1234", "super_admin"),
        ],
        [1.6, 2.2, 1.2, 1.4],
    )
    heading(doc, "A.1 Struktur Folder Penting", 2)
    matrix(
        doc,
        ["Path", "Fungsi"],
        [
            ("client/lib", "Source Flutter untuk UI."),
            ("client/lib/main_user.dart", "Entrypoint APK User."),
            ("client/lib/main_counseling.dart", "Entrypoint APK Counseling."),
            ("client/lib/main.dart", "Entrypoint Web Admin."),
            ("client/android/app/build.gradle.kts", "Product flavor dan konfigurasi Android."),
            ("server/src", "Source backend API."),
            ("server/data/nsd.json", "Data seed/demo berbasis file."),
            ("firebase.json", "Konfigurasi Firebase Hosting."),
        ],
        [2.8, 3.4],
    )

    heading(doc, "Lampiran B Command Build dan Deploy", 1)
    heading(doc, "B.1 Typecheck", 2)
    matrix(doc, ["Perintah", "Tujuan"], [("npm run typecheck", "Menjalankan TypeScript typecheck dan flutter analyze.")], [2.8, 3.4])
    heading(doc, "B.2 Build APK", 2)
    matrix(
        doc,
        ["Target", "Perintah", "Output"],
        [
            ("APK User", "npm run build:apk:user", "client/build/app/outputs/flutter-apk/app-user-release.apk"),
            ("APK Counseling", "npm run build:apk:counseling", "client/build/app/outputs/flutter-apk/app-counseling-release.apk"),
        ],
        [1.5, 2.4, 2.5],
    )
    heading(doc, "B.3 Install APK ke HP", 2)
    matrix(
        doc,
        ["Target", "Perintah"],
        [
            ("User", "adb install -r client/build/app/outputs/flutter-apk/app-user-release.apk"),
            ("Counseling", "adb install -r client/build/app/outputs/flutter-apk/app-counseling-release.apk"),
        ],
        [1.5, 4.9],
    )
    heading(doc, "B.4 Deploy Firebase Hosting", 2)
    steps(
        doc,
        [
            "npm run build",
            "firebase projects:list",
            "firebase deploy --only hosting",
            "Buka Hosting URL dan lakukan hard refresh apabila browser masih membaca build lama.",
        ],
    )
    heading(doc, "B.5 Deploy Railway", 2)
    steps(
        doc,
        [
            "Pastikan service Railway terhubung ke repository GitHub.",
            "Isi environment variables: CLIENT_ORIGIN, JWT_SECRET, FIREBASE_PROJECT_ID, FIREBASE_SERVICE_ACCOUNT_JSON, PORT.",
            "Push commit ke GitHub.",
            "Cek Deploy Logs hingga status Active.",
            "Buka endpoint /api/health pada domain Railway.",
        ],
    )

    heading(doc, "Penutup", 1)
    para(
        doc,
        "Manual ini disusun sebagai panduan lengkap untuk menjalankan NSD sebagai sistem donasi dengan tiga permukaan utama: APK User, APK Counseling, dan Web Admin Panel. Apabila ada perubahan desain atau workflow, bagian screenshot dan SOP perlu diperbarui agar dokumen tetap sesuai dengan aplikasi yang sedang berjalan.",
    )
    callout(
        doc,
        "Versi dokumen",
        "Dokumen ini dibuat dari source code dan screenshot build lokal pada 08 Juni 2026. File akhir: USER_MANUAL_NSD_LENGKAP_BAB_SCREENSHOT.docx.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
